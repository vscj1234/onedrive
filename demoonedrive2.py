import msal
from flask import Flask, request, redirect, session, url_for, jsonify, render_template
import requests
import os
from io import BytesIO
from docx import Document
import chromadb
from chromadb.utils import embedding_functions
from typing import List, Tuple
from dotenv import load_dotenv
import logging
from openai import OpenAI
import zipfile
import json
import hashlib
import mimetypes
from werkzeug.utils import secure_filename
from datetime import datetime
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.embeddings import OpenAIEmbeddings
from langchain.prompts import PromptTemplate
from langchain.llms import OpenAI as LangchainOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.llms import OpenAI as LangchainOpenAI


app = Flask(__name__)
app.secret_key = os.urandom(24)  # Set this to a secure random string

# Load environment variables
load_dotenv()

# Azure AD configuration
CLIENT_ID = "b88e7b71-f38b-4e8a-829c-28374df3e244"
CLIENT_SECRET = "Tme8Q~_zYbF2F.s.OdWpaC8cMi.gu28Qlx-m3aSD"
TENANT_ID = "e820bd4d-57c3-4334-94e4-c56a843a8cdc"
AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}"
REDIRECT_URI = 'http://localhost:5000/getAToken'
SCOPE = ['User.Read', 'Files.Read','Files.ReadWrite']

# OpenAI configuration
client = OpenAI(api_key=" ")

msal_app = msal.ConfidentialClientApplication(
    CLIENT_ID, authority=AUTHORITY,
    client_credential=CLIENT_SECRET
)

# Initialize Chroma
chroma_client = chromadb.Client()
collection = chroma_client.create_collection(name="onedrive_docs")

# Initialize Langchain components
embeddings = OpenAIEmbeddings()
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
vectorstore = Chroma(embedding_function=embeddings, persist_directory="./chromaonedrive_db")




logging.basicConfig(level=logging.INFO)

def read_docx(file_path_or_bytes) -> str:
    """
    Read content from a .docx file
    
    Args:
        file_path_or_bytes: Either a file path or bytes/BytesIO object containing the document
        
    Returns:
        str: Extracted text content from the document
    """
    try:
        # If input is bytes, convert to BytesIO
        if isinstance(file_path_or_bytes, bytes):
            file_stream = BytesIO(file_path_or_bytes)
        elif isinstance(file_path_or_bytes, BytesIO):
            file_stream = file_path_or_bytes
        else:
            file_stream = file_path_or_bytes
            
        # Ensure we're at the start of the stream
        file_stream.seek(0)
        
        doc = Document(file_stream)
        full_text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Only add non-empty cells
                        full_text.append(cell.text.strip())
        
        return " ".join(full_text)
        
    except ValueError as e:
        logging.error(f"Error reading DOCX file: {str(e)}")
        raise ValueError(f"Invalid document format: {str(e)}")
    except zipfile.BadZipFile:
        logging.error("File is not a valid ZIP file (possibly corrupted Word document)")
        raise ValueError("Corrupted or invalid Word document")
    except Exception as e:
        logging.error(f"Unexpected error reading DOCX: {str(e)}")
        raise


def fetch_file_content(file_id: str, access_token: str) -> Tuple[str, bytes, str]:
    headers = {'Authorization': f'Bearer {access_token}'}
    file_data = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}', headers=headers).json()
    file_name = file_data['name']
    file_type = file_data.get('file', {}).get('mimeType', '')
    
    content_response = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}/content', headers=headers)
    return file_name, content_response.content, file_type

def get_file_metadata(file_id, access_token):
    headers = {'Authorization': f'Bearer {access_token}'}
    file_data = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}', headers=headers).json()
    return {
        'id': file_id,
        'name': file_data['name'],
        'lastModifiedDateTime': file_data['lastModifiedDateTime']
    }

def load_processed_files():
    if os.path.exists('processed_files.json'):
        with open('processed_files.json', 'r') as f:
            return json.load(f)
    return {}

def save_processed_files(processed_files):
    with open('processed_files.json', 'w') as f:
        json.dump(processed_files, f)

class FileProcessingCache:
    def __init__(self, cache_dir="./cache"):
        self.cache_dir = cache_dir
        self.cache_file = os.path.join(cache_dir, "file_cache.json")
        self.embeddings_dir = os.path.join(cache_dir, "embeddings")
        os.makedirs(self.cache_dir, exist_ok=True)
        os.makedirs(self.embeddings_dir, exist_ok=True)
        self.cache = self._load_cache()

    def _load_cache(self):
        if os.path.exists(self.cache_file):
            with open(self.cache_file, 'r') as f:
                return json.load(f)
        return {}

    def _save_cache(self):
        with open(self.cache_file, 'w') as f:
            json.dump(self.cache, f)

    def _calculate_file_hash(self, file_content: bytes) -> str:
        return hashlib.md5(file_content).hexdigest()

    def needs_processing(self, file_id: str, file_content: bytes, last_modified: str) -> bool:
        file_hash = self._calculate_file_hash(file_content)
        
        if file_id not in self.cache:
            return True
            
        cached_info = self.cache[file_id]
        return (cached_info['hash'] != file_hash or 
                cached_info['last_modified'] != last_modified)

    def update_cache(self, file_id: str, file_content: bytes, metadata: dict):
        self.cache[file_id] = {
            'hash': self._calculate_file_hash(file_content),
            'last_modified': metadata['lastModifiedDateTime'],
            'processed_date': datetime.now().isoformat(),
            'metadata': metadata
        }
        self._save_cache()

def process_files_optimized(access_token):
    cache = FileProcessingCache()
    headers = {'Authorization': f'Bearer {access_token}'}
    processed_count = 0
    skipped_count = 0

    def process_item_optimized(item, parent_path=""):
        nonlocal processed_count, skipped_count
        item_id = item['id']
        item_metadata = get_file_metadata(item_id, access_token)
        item_path = os.path.join(parent_path, item_metadata['name'])

        if 'folder' in item:
            folder_content = requests.get(
                f'https://graph.microsoft.com/v1.0/me/drive/items/{item_id}/children',
                headers=headers
            ).json()
            for child_item in folder_content.get('value', []):
                process_item_optimized(child_item, item_path)
        else:
            try:
                file_name, file_content, file_type = fetch_file_content(item_id, access_token)
                
                if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    if not cache.needs_processing(item_id, file_content, item_metadata['lastModifiedDateTime']):
                        logging.info(f"Skipping unchanged file: {item_path}")
                        skipped_count += 1
                        return

                    text_content = read_docx(file_content)
                    if text_content:
                        chunks = text_splitter.split_text(text_content)
                        
                        metadatas = [{
                            "source": item_path,
                            "file_name": file_name,
                            "file_id": item_id,
                            "chunk_index": i,
                            "processing_date": datetime.now().isoformat()
                        } for i in range(len(chunks))]
                        
                        vectorstore.add_texts(
                            texts=chunks,
                            metadatas=metadatas,
                            ids=[f"{item_id}_chunk_{i}" for i in range(len(chunks))]
                        )
                        
                        cache.update_cache(item_id, file_content, item_metadata)
                        processed_count += 1
                        logging.info(f"Successfully processed file: {item_path}")
                    else:
                        logging.warning(f"No content extracted from file: {item_path}")
            except Exception as e:
                logging.error(f"Error processing file {item_path}: {str(e)}")
                raise

    try:
        root_items = requests.get('https://graph.microsoft.com/v1.0/me/drive/root/children', headers=headers).json()
        for item in root_items.get('value', []):
            process_item_optimized(item)
        
        logging.info(f"Processing complete. Processed: {processed_count} files, Skipped: {skipped_count} files")
        return processed_count, skipped_count
    except Exception as e:
        logging.error(f"Error in file processing: {str(e)}")
        raise


# Initialize ConversationalRetrievalChain
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
qa_chain = ConversationalRetrievalChain.from_llm(
    LangchainOpenAI(temperature=0.7, max_tokens=150),
    vectorstore.as_retriever(search_kwargs={"k": 3}),
    memory=memory
)


def query_documents(query: str, n_results: int = 3) -> List[Tuple[str, str]]:
    results = collection.query(
        query_texts=[query],
        n_results=n_results
    )
    return list(zip(results['documents'][0], results['metadatas'][0]))

def generate_response(query: str, chat_history: List[dict] = None) -> Tuple[str, List[dict]]:
    if chat_history is None:
        chat_history = []
    try:
        result = qa_chain({"question": query, "chat_history": chat_history})
        answer = result["answer"]
        
        # Retrieve the most relevant documents
        relevant_docs = vectorstore.similarity_search(query, k=3)
        doc_references = []
        seen_files = set()
        
        for doc in relevant_docs:
            file_id = doc.metadata.get("file_id")
            if file_id and file_id not in seen_files:
                doc_references.append({
                    "file_name": doc.metadata.get("file_name", "Unknown"),
                    "source": doc.metadata.get("source", "Unknown path"),
                    "file_id": file_id
                })
                seen_files.add(file_id)
        
        logging.info(f"Generated response for query: {query}")
        return answer, doc_references
    except Exception as e:
        logging.error(f"Error generating response: {str(e)}")
        return "I'm sorry, but I encountered an error while processing your request. Please try again later.", []

@app.route('/')
def index():
    if not session.get("user"):
        return redirect(url_for('login'))

    return render_template('index.html')

@app.route('/open_document/<file_id>')
def open_document(file_id):
    logging.info(f"Attempting to open document with ID: {file_id}")
    
    # 1. Check authentication
    if not session.get('token'):
        logging.error("No authentication token found in session")
        return jsonify({"error": "Not authenticated"}), 401

    # 2. Validate file_id
    if not file_id or file_id == "Unknown":
        logging.error(f"Invalid file_id received: {file_id}")
        return jsonify({"error": "Invalid file ID"}), 400

    access_token = session['token']
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }

    try:
        # 3. First attempt: Get direct download URL
        logging.info(f"Attempting to get direct download URL for file: {file_id}")
        response = requests.get(
            f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}',
            headers=headers
        )
        
        logging.info(f"Direct URL Response Status: {response.status_code}")
        logging.debug(f"Direct URL Response Content: {response.text}")

        if response.status_code == 200:
            file_data = response.json()
            download_url = file_data.get('@microsoft.graph.downloadUrl')
            
            if download_url:
                logging.info("Successfully retrieved download URL")
                return jsonify({"download_url": download_url})

        # 4. Second attempt: Create sharing link
        logging.info("Attempting to create sharing link")
        sharing_body = {
            "type": "view",
            "scope": "organization"  # Changed from anonymous to organization
        }

        sharing_response = requests.post(
            f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}/createLink',
            headers=headers,
            json=sharing_body
        )
        
        logging.info(f"Sharing Response Status: {sharing_response.status_code}")
        logging.debug(f"Sharing Response Content: {sharing_response.text}")

        if sharing_response.status_code == 200:
            share_data = sharing_response.json()
            share_link = share_data.get('link', {}).get('webUrl')
            
            if share_link:
                logging.info("Successfully created sharing link")
                return jsonify({"download_url": share_link})
            else:
                logging.error("No webUrl in sharing response")
                return jsonify({"error": "No sharing URL in response"}), 500

        # 5. If both methods fail, try to get a temporary access URL
        logging.info("Attempting to get temporary access URL")
        temp_url_response = requests.get(
            f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}?select=id,@microsoft.graph.downloadUrl',
            headers=headers
        )
        
        if temp_url_response.status_code == 200:
            temp_url_data = temp_url_response.json()
            temp_download_url = temp_url_data.get('@microsoft.graph.downloadUrl')
            
            if temp_download_url:
                logging.info("Successfully retrieved temporary download URL")
                return jsonify({"download_url": temp_download_url})

        # 6. If all attempts fail, return detailed error
        error_msg = {
            "error": "Unable to generate access URL",
            "details": {
                "direct_url_status": response.status_code,
                "sharing_status": sharing_response.status_code,
                "file_id": file_id
            }
        }
        logging.error(f"All attempts to get URL failed: {error_msg}")
        return jsonify(error_msg), 500

    except requests.exceptions.RequestException as e:
        logging.error(f"Network error while accessing file: {str(e)}")
        return jsonify({"error": "Network error while accessing file"}), 500
    except Exception as e:
        logging.error(f"Unexpected error while accessing file: {str(e)}")
        return jsonify({"error": "Unexpected error while accessing file"}), 500

# Add this helper function to verify file accessibility
def verify_file_access(file_id, access_token):
    """Verify if the file is accessible and return its metadata"""
    headers = {'Authorization': f'Bearer {access_token}'}
    
    try:
        response = requests.get(
            f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}',
            headers=headers
        )
        if response.status_code == 200:
            return True, response.json()
        return False, response.json()
    except Exception as e:
        return False, str(e)

@app.route('/login')
def login():
    auth_url = msal_app.get_authorization_request_url(SCOPE, redirect_uri=REDIRECT_URI)
    return redirect(auth_url)

# In your main Flask app
@app.route('/getAToken')
def authorized():
    if request.args.get('code'):
        token_result = msal_app.acquire_token_by_authorization_code(
            request.args['code'], scopes=SCOPE, redirect_uri=REDIRECT_URI)
        if "access_token" in token_result:
            session['user'] = token_result.get('id_token_claims')
            session['token'] = token_result.get('access_token')
            try:
                processed, skipped = process_files_optimized(session['token'])
                logging.info(f"File processing complete: {processed} processed, {skipped} skipped")
            except Exception as e:
                logging.error(f"Error processing files: {str(e)}")
    return redirect(url_for('index'))

@app.route('/upload', methods=['POST'])
def upload_file():
    # 1. Authentication Check
    if not session.get('token'):
        logging.error("Upload attempted without valid token")
        return jsonify({'error': 'Authentication required'}), 403
    
    # 2. File Validation
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    try:
        # 3. File Processing
        filename = secure_filename(file.filename)
        
        # Read file into BytesIO object
        file_stream = BytesIO()
        file.save(file_stream)
        file_stream.seek(0)  # Reset stream position to beginning
        
        # Get file size
        file_size = file_stream.getbuffer().nbytes
        
        if file_size == 0:
            return jsonify({'error': 'Empty file'}), 400
            
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return jsonify({'error': 'Only .docx files are supported'}), 400

        # 4. Graph API Headers
        headers = {
            'Authorization': f'Bearer {session["token"]}',
            'Content-Type': 'application/json'
        }
        
        # 5. Create Upload Session
        upload_session_url = 'https://graph.microsoft.com/v1.0/me/drive/root:/Documents/' + filename + ':/createUploadSession'
        
        upload_session_data = {
            "item": {
                "@microsoft.graph.conflictBehavior": "rename",
                "name": filename
            }
        }
        
        try:
            # Create upload session with proper error handling
            upload_session_response = requests.post(
                upload_session_url,
                headers=headers,
                json=upload_session_data
            )
            
            if not upload_session_response.ok:
                logging.error(f"Upload session creation failed: {upload_session_response.text}")
                if upload_session_response.status_code == 401:
                    return jsonify({'error': 'Token expired. Please refresh your session.'}), 401
                elif upload_session_response.status_code == 403:
                    return jsonify({'error': 'Permission denied. Please check your OneDrive permissions.'}), 403
                else:
                    return jsonify({
                        'error': 'Failed to create upload session',
                        'details': upload_session_response.text
                    }), upload_session_response.status_code
            
            upload_url = upload_session_response.json().get('uploadUrl')
            if not upload_url:
                return jsonify({'error': 'No upload URL received'}), 500

            # 6. Upload File Content
            upload_headers = {
                'Content-Length': str(file_size),
                'Content-Range': f'bytes 0-{file_size-1}/{file_size}'
            }
            
            # Reset stream position before upload
            file_stream.seek(0)
            
            upload_response = requests.put(
                upload_url,
                headers=upload_headers,
                data=file_stream
            )
            
            if not upload_response.ok:
                logging.error(f"File upload failed: {upload_response.text}")
                return jsonify({
                    'error': 'Failed to upload file',
                    'details': upload_response.text
                }), upload_response.status_code

            # 7. Process Uploaded File
            file_data = upload_response.json()
            file_id = file_data.get('id')
            
            if not file_id:
                return jsonify({'error': 'No file ID received'}), 500

            # 8. Process Document Content
            try:
                # Reset stream position before reading content
                file_stream.seek(0)
                
                # Extract text content
                text_content = read_docx(file_stream)
                if not text_content:
                    return jsonify({'error': 'No content could be extracted from file'}), 400

                chunks = text_splitter.split_text(text_content)
                
                metadatas = [{
                    "source": f"Documents/{filename}",
                    "file_name": filename,
                    "file_id": file_id,
                    "chunk_index": i,
                    "upload_date": datetime.now().isoformat()
                } for i in range(len(chunks))]
                
                vectorstore.add_texts(
                    texts=chunks,
                    metadatas=metadatas,
                    ids=[f"{file_id}_chunk_{i}" for i in range(len(chunks))]
                )
                
                return jsonify({
                    'success': True,
                    'message': 'File uploaded and processed successfully',
                    'file_id': file_id,
                    'file_name': filename,
                    'chunks_processed': len(chunks)
                })
                
            except ValueError as ve:
                logging.error(f"Error processing document content: {str(ve)}")
                return jsonify({
                    'error': 'Error processing document content',
                    'details': str(ve)
                }), 400
                
            except Exception as e:
                logging.error(f"Error processing document content: {str(e)}")
                return jsonify({
                    'error': 'Error processing document content',
                    'details': str(e)
                }), 500
                
        except requests.exceptions.RequestException as e:
            logging.error(f"Network error during upload: {str(e)}")
            return jsonify({
                'error': 'Network error during upload',
                'details': str(e)
            }), 500
            
    except Exception as e:
        logging.error(f"Unexpected error in upload: {str(e)}")
        return jsonify({
            'error': 'Upload failed',
            'details': str(e)
        }), 500

# Add status check endpoint
@app.route('/upload/status/<file_id>')
def check_upload_status(file_id):
    if not session.get('token'):
        return jsonify({'error': 'Not authenticated'}), 401
        
    try:
        # Verify file exists and is processed
        headers = {'Authorization': f'Bearer {session["token"]}'}
        response = requests.get(
            f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}',
            headers=headers
        )
        
        if response.status_code == 200:
            return jsonify({
                'status': 'complete',
                'file_info': response.json()
            })
        return jsonify({'status': 'pending'})
        
    except Exception as e:
        logging.error(f"Error checking upload status: {str(e)}")
        return jsonify({'error': 'Status check failed'}), 500

@app.route('/chat', methods=['POST'])
def chat():
    if not session.get('token'):
        logging.warning("Unauthorized access attempt to chat endpoint")
        return jsonify({"error": "Not authenticated"}), 401
    
    query = request.json.get('query')
    if not query:
        logging.warning("Empty query received")
        return jsonify({"error": "No query provided"}), 400
    
    chat_history = session.get('chat_history', [])
    
    try:
        logging.info(f"Processing query: {query}")
        
        # Use the qa_chain to generate a response
        result = qa_chain({
            "question": query, 
            "chat_history": chat_history
        })
        response = result["answer"]
        
        # Retrieve and process the most relevant documents
        relevant_docs = vectorstore.similarity_search(query, k=3)
        
        # Create a set to track unique file IDs
        seen_file_ids = set()
        doc_references = []
        
        for doc in relevant_docs:
            file_id = doc.metadata.get("file_id")
            if file_id and file_id != "Unknown" and file_id not in seen_file_ids:
                doc_references.append({
                    "file_name": doc.metadata.get("file_name", "Unknown"),
                    "source": doc.metadata.get("source", "Unknown path"),
                    "file_id": file_id,
                    "relevance_score": doc.metadata.get("relevance_score", 0.0)
                })
                seen_file_ids.add(file_id)
                
                # Log successful document reference
                logging.info(f"Found relevant document: {doc.metadata.get('file_name')} (ID: {file_id})")
        
        # Sort documents by relevance score if available
        doc_references.sort(key=lambda x: x.get("relevance_score", 0), reverse=True)
        
        # Update chat history with timestamped entries
        timestamp = datetime.now().isoformat()
        chat_history.append({
            "role": "user",
            "content": query,
            "timestamp": timestamp
        })
        chat_history.append({
            "role": "assistant",
            "content": response,
            "timestamp": timestamp,
            "documents": [doc["file_name"] for doc in doc_references]
        })
        
        # Keep only the last 10 messages in chat history
        session['chat_history'] = chat_history[-10:]
        
        # Prepare the response
        api_response = {
            "response": response,
            "relevant_documents": doc_references,
            "timestamp": timestamp
        }
        
        # Log successful response
        logging.info(f"Successfully processed query with {len(doc_references)} relevant documents")
        
        return jsonify(api_response)
    
    except Exception as e:
        # Log the full error with traceback
        logging.error(f"Error in chat route: {str(e)}", exc_info=True)
        
        # Check if it's a specific type of error we can handle
        if isinstance(e, (ValueError, KeyError)):
            error_message = "Invalid input or missing data"
            status_code = 400
        else:
            error_message = "An internal error occurred while processing your request"
            status_code = 500
        
        return jsonify({
            "error": error_message,
            "error_type": type(e).__name__,
            "timestamp": datetime.now().isoformat()
        }), status_code


if __name__ == "__main__":
    app.run(debug=True)